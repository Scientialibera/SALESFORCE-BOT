# This script generates a Word document (DOCX) with the cost breakdown the user requested.
# If python-docx isn't available, it falls back to an RTF file that opens in Word.

from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime

# -----------------------
# Assumptions (from user)
# -----------------------
# Model pricing (user-provided; see doc narrative for caveat to verify against pricing page)
gpt41_input_per_1k   = Decimal("0.003")   # $ per 1k input tokens  => $3.00 / 1M
gpt41_output_per_1k  = Decimal("0.012")   # $ per 1k output tokens => $12.00 / 1M
emb_price_per_1m     = Decimal("0.130")   # $ per 1M tokens for text-embedding-3-large

# First pass scale
accounts_first_pass  = 100_000
judge_input_tokens_per_acct  = 2_000
judge_output_tokens_per_acct = 50

# Embeddings – "max 5M tokens for everything"
emb_first_pass_tokens = 5_000_000

# Monthly run scale
new_accounts_per_month = 200          # "running costs ... 200 max new accounts per month"
reruns_accounts_max    = 1_000        # "max running reruns ... 1000 accounts MAX" per month

# ---------------
# Helper functions
# ---------------
def money(x: Decimal) -> str:
    return f"${x.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP):,}"

def per_1k_cost(tokens: int, rate_per_1k: Decimal) -> Decimal:
    # tokens / 1000 * rate
    return (Decimal(tokens) / Decimal(1000)) * rate_per_1k

def per_1m_cost(tokens: int, rate_per_1m: Decimal) -> Decimal:
    return (Decimal(tokens) / Decimal(1_000_000)) * rate_per_1m

# -------------------------
# First pass LLM (GPT-4.1)
# -------------------------
first_input_tokens  = accounts_first_pass * judge_input_tokens_per_acct
first_output_tokens = accounts_first_pass * judge_output_tokens_per_acct

first_input_cost  = per_1k_cost(first_input_tokens, gpt41_input_per_1k)
first_output_cost = per_1k_cost(first_output_tokens, gpt41_output_per_1k)
first_total_cost  = first_input_cost + first_output_cost

# ----------------------
# First pass Embeddings
# ----------------------
first_embeddings_cost = per_1m_cost(emb_first_pass_tokens, emb_price_per_1m)

# --------------------
# Monthly GPT-4.1 LLM
# --------------------
monthly_input_tokens  = (new_accounts_per_month + reruns_accounts_max) * judge_input_tokens_per_acct
monthly_output_tokens = (new_accounts_per_month + reruns_accounts_max) * judge_output_tokens_per_acct

monthly_input_cost  = per_1k_cost(monthly_input_tokens, gpt41_input_per_1k)
monthly_output_cost = per_1k_cost(monthly_output_tokens, gpt41_output_per_1k)
monthly_total_cost  = monthly_input_cost + monthly_output_cost

annual_recurring_total = monthly_total_cost * Decimal(12)

# -----------------
# Build the document
# -----------------
title = "OpenAI Cost Breakdown — GPT‑4.1 & Embeddings"
subtitle = f"Prepared {datetime.now().strftime('%B %d, %Y')}"

# Data blocks for tables
rates_table = [
    ["Item", "Rate"],
    ["GPT‑4.1 Input",  "$0.003 per 1K tokens  (=$3.00 / 1M)"],
    ["GPT‑4.1 Output", "$0.012 per 1K tokens (=$12.00 / 1M)"],
    ["text-embedding-3-large", "$0.130 per 1M tokens"],
]

first_pass_table = [
    ["Category", "Tokens", "Unit Rate", "Cost"],
    ["LLM Input",  f"{first_input_tokens:,}",  "$0.003 / 1K", money(first_input_cost)],
    ["LLM Output", f"{first_output_tokens:,}", "$0.012 / 1K", money(first_output_cost)],
    ["— LLM Total —", "", "", money(first_total_cost)],
    ["Embeddings (1x)", f"{emb_first_pass_tokens:,}", "$0.130 / 1M", money(first_embeddings_cost)],
    ["— One‑time First Pass Total —", "", "", money(first_total_cost + first_embeddings_cost)],
]

monthly_table = [
    ["Category", "Tokens / mo", "Unit Rate", "Monthly Cost"],
    ["LLM Input",  f"{monthly_input_tokens:,}",  "$0.003 / 1K", money(monthly_input_cost)],
    ["LLM Output", f"{monthly_output_tokens:,}", "$0.012 / 1K", money(monthly_output_cost)],
    ["— Monthly LLM Total —", "", "", money(monthly_total_cost)],
    ["— Annualized Recurring (12×) —", "", "", money(annual_recurring_total)],
]

assumptions_list = [
    "LLM judge size = 2,000 input tokens + 50 output tokens per account.",
    "Initial pass over 100,000 accounts.",
    "Embeddings one‑time volume capped at 5,000,000 tokens (text-embedding-3-large).",
    "Monthly operations: up to 200 new accounts and up to 1,000 reruns.",
    "Pricing uses client‑provided rates for GPT‑4.1 ($0.003/1K in, $0.012/1K out) and OpenAI’s published rate for text‑embedding‑3‑large ($0.13/1M tokens).",
]

# Try to build a DOCX; fall back to RTF if python-docx isn't installed
docx_path = "OpenAI_Cost_Breakdown.docx"
rtf_path  = "OpenAI_Cost_Breakdown.rtf"

made_docx = False
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    document = Document()

    # Title
    p = document.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(20)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = document.add_paragraph(subtitle)
    p.runs[0].font.size = Pt(11)
    p.runs[0].italic = True

    # Section: Assumptions
    document.add_paragraph().add_run("Assumptions").bold = True
    for a in assumptions_list:
        document.add_paragraph(a, style="List Bullet")

    # Section: Rates
    document.add_paragraph().add_run("Rates").bold = True
    table = document.add_table(rows=1, cols=len(rates_table[0]))
    hdr_cells = table.rows[0].cells
    for i, head in enumerate(rates_table[0]):
        hdr_cells[i].text = head
    for row in rates_table[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    # Section: First Pass
    document.add_paragraph().add_run("One‑time First Pass").bold = True
    table = document.add_table(rows=1, cols=len(first_pass_table[0]))
    hdr_cells = table.rows[0].cells
    for i, head in enumerate(first_pass_table[0]):
        hdr_cells[i].text = head
    for row in first_pass_table[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    # Section: Monthly Recurring
    document.add_paragraph().add_run("Monthly Recurring (plus Annualized)").bold = True
    table = document.add_table(rows=1, cols=len(monthly_table[0]))
    hdr_cells = table.rows[0].cells
    for i, head in enumerate(monthly_table[0]):
        hdr_cells[i].text = head
    for row in monthly_table[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    # Notes
    document.add_paragraph().add_run("Notes").bold = True
    notes = [
        "Costs exclude network, storage, orchestration, and non-OpenAI compute.",
        "Actual pricing may vary by region and contract; confirm against OpenAI’s pricing page before committing budgets.",
    ]
    for n in notes:
        document.add_paragraph(n, style="List Bullet")

    document.save(docx_path)
    made_docx = True
except Exception as e:
    made_docx = False
    err = str(e)

# Fallback to RTF if needed
if not made_docx:
    def rtf_escape(s: str) -> str:
        return s.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")

    def rtf_para(text: str, bold=False, italic=False):
        prefix = ""
        if bold: prefix += "\\b "
        if italic: prefix += "\\i "
        suffix = ""
        if bold: suffix += "\\b0 "
        if italic: suffix += "\\i0 "
        return f"\\par {prefix}{rtf_escape(text)}{suffix}"

    lines = []
    lines.append(r"{\rtf1\ansi")
    lines.append(rtf_para(title, bold=True))
    lines.append(rtf_para(subtitle, italic=True))
    lines.append(rtf_para(""))
    lines.append(rtf_para("Assumptions", bold=True))
    for a in assumptions_list:
        lines.append(rtf_para(f"• {a}"))
    lines.append(rtf_para(""))
    lines.append(rtf_para("Rates", bold=True))
    for r in rates_table[1:]:
        lines.append(rtf_para(f"- {r[0]}: {r[1]}"))
    lines.append(rtf_para(""))
    lines.append(rtf_para("One‑time First Pass", bold=True))
    for r in first_pass_table[1:]:
        label = r[0]
        tokens = r[1]
        rate = r[2]
        cost = r[3]
        if "—" in label:
            lines.append(rtf_para(f"{label}  {cost}", bold=True))
        else:
            lines.append(rtf_para(f"{label}: tokens={tokens}, rate={rate}, cost={cost}"))
    lines.append(rtf_para(""))
    lines.append(rtf_para("Monthly Recurring (plus Annualized)", bold=True))
    for r in monthly_table[1:]:
        label = r[0]
        tokens = r[1]
        rate = r[2]
        cost = r[3]
        if "—" in label:
            lines.append(rtf_para(f"{label}  {cost}", bold=True))
        else:
            lines.append(rtf_para(f"{label}: tokens={tokens}, rate={rate}, monthly cost={cost}"))
    lines.append(rtf_para(""))
    lines.append(rtf_para("Notes", bold=True))
    lines.append(rtf_para("• Costs exclude network, storage, orchestration, and non-OpenAI compute."))
    lines.append(rtf_para("• Actual pricing may vary by region and contract; confirm against OpenAI’s pricing page before committing budgets."))
    lines.append("}")

    with open(rtf_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

{"result_path": "OpenAI_Cost_Breakdown.docx" if made_docx else "OpenAI_Cost_Breakdown.rtf", "made_docx": made_docx, "monthly_total_cost": str(monthly_total_cost), "annual_total": str(annual_recurring_total), "first_total_cost": str(first_total_cost), "first_embeddings_cost": str(first_embeddings_cost)}
